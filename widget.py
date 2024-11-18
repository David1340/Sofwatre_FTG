# This Python file uses the following encoding: utf-8
import sys

from PySide6.QtWidgets import QApplication, QWidget

# Important:
# You need to run the following command to generate the ui_form.py file
#     pyside6-uic form.ui -o ui_form.py, or
#     pyside2-uic form.ui -o ui_form.py
from ui_form import Ui_Widget

from openpyxl import Workbook,load_workbook #Para trabalhar com planilha
from datetime import datetime,date #Para acessar data e horário
from docx import Document #Para manipular arquivos .Docx
#from docx.shared import Pt
#from docx.enum.text import WD_ALIGN_PARAGRAPH
import sqlite3
import serial
import serial.tools.list_ports
import time
import matplotlib.pyplot as plt
from docx.shared import Inches
from docx.enum.shape import WD_INLINE_SHAPE


class Experimento():
    def __init__(self,nome,id):
        self.data = date.today().strftime("%d/%m/%Y")
        self.cont_corridas = 0
        self.exercicios = {
            "Corrida": [0] * 8,
            "Burpee": 0,
            "Lunge": 0,
            "Farmer_Walk": 0,
            "Abdominal": 0,
            "Jump_Squat": 0,
            "Medball_Alto": 0,
            "Medball_Solo": 0,
            "Terra_Remada": 0,

        }
        self.ID_Pessoa = 0
        self.nome = nome
        self.update_ID_Pessoa()
        self._id = id




    def __str__(self):
            return f"""
            Nome: {self.nome}
            ID: {self._id}
            Data: {self.data}
            Contagem Corridas: {self.cont_corridas}
            ID Pessoa: {self.ID_Pessoa}
            Exercícios: {self.exercicios}
            """

    def update_corridas(self,value):
        print(self.exercicios["Corrida"][self.cont_corridas])
        self.exercicios["Corrida"][self.cont_corridas] = value
        self.cont_corridas += 1

    def update_ID_Pessoa(self):
        # Conectar ao banco de dados
        conexao = sqlite3.connect('banco de dados.db')
        cursor = conexao.cursor()

        cursor.execute("SELECT ID_Pessoa FROM pessoas WHERE nome = ?", (self.nome,))

        # Obter o resultado da consulta
        resultados = cursor.fetchone()
        self.ID_Pessoa = resultados[0]
        # Fechar a conexão com o banco de dados
        conexao.close()


class Widget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.ui = Ui_Widget()
        self.ui.setupUi(self)
        self.ui.PushButtonGerarRelatorio.clicked.connect(self.gerar_relatorio)
        self.ui.PushButtonGerarPlanilha.clicked.connect(self.gerar_planilha)
        self.ui.PushButtonTotem.clicked.connect(self.solicitar_informacoes_totem)
        self.ui.PushButtonLimparAtualizacoes.clicked.connect(self.Limpar_atualizacoes_totens)
        self.ui.PushButtonCadastro.clicked.connect(self.cadastro)
        self.atualizar_nomes_participantes()
        self.ui.pushButtonUSB2.clicked.connect(self.atualizar_portas)
        self.ui.pushButtonUSB.clicked.connect(self.conectar_usb)
        self.serial_port = None
        self.experimentos = []
        self.criar_experimentos()
        self.dicionario_exercicios = {}
        self.criar_dicionario_exercicios()
        self.ui.pushButtonBancoDados.clicked.connect(self.enviar_informacoes_banco_dados)


    def enviar_informacoes_banco_dados(self):
        # Itera sobre os dados em 'data.txt'
        time_intervals = {}
        with open('log_serial.txt') as arquivo:
            linhas = arquivo.readlines()
            for linha in linhas:
                entry = linha.split(",")
                id_totem = int(entry[0])  # IDTOTEM
                id_value = int(entry[1])  # ID
                btn = int(entry[2])       # BTN
                timestamp = int(entry[3]) # TIMESTAMP

                # Inicializa o IDTOTEM no dicionário se ainda não existe
                if id_totem not in time_intervals:
                    time_intervals[id_totem] = {}

                # Inicializa o ID no dicionário de IDTOTEM se ainda não existe
                if id_value not in time_intervals[id_totem]:
                    time_intervals[id_totem][id_value] = []

                # Se BTN=1, armazena como tempo inicial; se BTN=2, armazena como tempo final
                if btn == 1:
                    # Adiciona uma nova entrada para o tempo inicial sem tempo final ainda
                    time_intervals[id_totem][id_value].append([timestamp, None])
                elif btn == 2:
                    # Atualiza o último tempo inicial com o tempo final correspondente
                    for interval in time_intervals[id_totem][id_value]:
                        if interval[1] is None:  # Encontra o tempo inicial sem um final correspondente
                            interval[1] = timestamp
                            break

            print(time_intervals)

            for id_totem in time_intervals:
                    for id_value in time_intervals[id_totem]:
                        time_intervals[id_totem][id_value] = [
                            tuple(interval) for interval in time_intervals[id_totem][id_value]
                        ]

            for id_totem, ids in time_intervals.items():
                    print(f"\nIDTOTEM: {id_totem}")
                    print(f"{'ID':<5}{'Tempo Inicial (TIMESTAMP)':<25}{'Tempo Final (TIMESTAMP)':<25}")
                    print("-" * 55)



                    # Exibe os intervalos de tempo para cada ID dentro do IDTOTEM
                    for id_value, intervals in ids.items():
                        for interval in intervals:
                            tempo_inicial = interval[0]
                            tempo_final = interval[1] if interval[1] is not None else "N/A"

                            for indice, exp in enumerate(self.experimentos):
                                if(id_value == exp._id):
                                    if(id_totem == self.dicionario_exercicios["Corrida"]):
                                        self.experimentos[indice].update_corridas(tempo_final - tempo_inicial)
                                    else:
                                        for nome_exercicio, tempos in exp.exercicios.items():
                                            if(id_totem == self.dicionario_exercicios[nome_exercicio]):
                                                self.experimentos[indice].exercicios[nome_exercicio] = tempo_final - tempo_inicial



                            # Imprime os valores de timestamp para o ID dentro do IDTOTEM
                            print(f"{id_value:<5}{tempo_inicial:<25}{tempo_final:<25}")

            for exp in self.experimentos:
                print(exp)

        conexao = sqlite3.connect('banco de dados.db')
        cursor = conexao.cursor()

        for exp in self.experimentos:
            # Conectando ao banco de dados


            # Obtendo o ID a partir do nome

            # Inserir dados diários
            tempos_exercios = [0]
            tempos_exercios = tempos_exercios*16

            #Tempo_total = Corrida + Burpee + Lunge + Farmer_Walk + Abdominal + Jump_Squat + Medball_Alto + \
            #Medball_Solo + Terra_remada

            cursor.execute("INSERT INTO experimentos (ID_Pessoa, Data, Burpee, Lunge, Farmer_Walk,\
            Abdominal, Jump_Squat, Medball_Alto, Medball_Solo, Terra_remada,Corrida1, Corrida2,\
            Corrida3,Corrida4,Corrida5,Corrida6,Corrida7,Corrida8) \
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)", \
            (exp.ID_Pessoa,exp.data,exp.exercicios["Burpee"],exp.exercicios["Lunge"],exp.exercicios["Farmer_Walk"],\
            exp.exercicios["Abdominal"],exp.exercicios["Jump_Squat"], exp.exercicios["Medball_Alto"],\
            exp.exercicios["Medball_Solo"],exp.exercicios["Terra_Remada"],exp.exercicios["Corrida"][0],\
            exp.exercicios["Corrida"][1],exp.exercicios["Corrida"][2],exp.exercicios["Corrida"][3],\
            exp.exercicios["Corrida"][4],exp.exercicios["Corrida"][5],exp.exercicios["Corrida"][6],\
            exp.exercicios["Corrida"][7]))


        # Confirmar a alteração
        conexao.commit()

        # Fechar a conexão
        conexao.close()

        with open('log_serial.txt', 'w') as file:
            pass  # Não faz nada, apenas limpa o conteúdo

    def cadastro(self):
        nome = self.ui.lineEditNome.text()
        idade = self.ui.lineEditIdade.text()
        if (not nome.strip()) and (not idade.strip()):  # strip() remove espaços em branco antes de checar
            self.ui.labelAvisoCadastro.setText("Todos os campos de texto estão vazios!")
        elif(nome == "cadastros"):
            # Conectando ao banco de dados
            conexao = sqlite3.connect('banco de dados.db')
            cursor = conexao.cursor()

            with open('cadastros.txt') as arquivo:
                linhas = arquivo.readlines()
                for linha in linhas:
                    dados = linha.split(",")
                    cursor.execute("INSERT INTO Pessoas (Nome, Idade) \
                    VALUES (?, ?)", \
                    (dados[0],dados[1]))

            # Confirmar a alteração
            conexao.commit()

            # Fechar a conexão
            conexao.close()

            self.ui.labelAvisoCadastro.setText("Os participantes foram cadastrados.")
            self.atualizar_nomes_participantes()

        else:
            if(idade == ""):
                idade = 0
            # Conectando ao banco de dados
            conexao = sqlite3.connect('banco de dados.db')
            cursor = conexao.cursor()

            cursor.execute("INSERT INTO Pessoas (Nome, Idade) \
            VALUES (?, ?)", \
            (nome,idade))

            # Confirmar a alteração
            conexao.commit()

            # Fechar a conexão
            conexao.close()

            self.ui.labelAvisoCadastro.setText("O participante foi cadastrado na base de dados.")
            self.atualizar_nomes_participantes()


    def gerar_relatorio(self):
        #Aquisição do nome do aprticipante escolhido
        nome = self.ui.ComboBoxNomes.currentText()

        # Conectar ao banco de dados
        conexao = sqlite3.connect('banco de dados.db')
        cursor = conexao.cursor()

        # Aquisição dos dados
        #data = date.today().strftime("%d/%m/%Y") caso queira na data de hoje sempre
        data = self.ui.calendarWidget.selectedDate().toString('dd/MM/yyyy')

        cursor.execute('''
        SELECT e.Data, e.Corrida1, e.Corrida2, e.Corrida3, e.Corrida4, e.Corrida5,\
        e.Corrida6, e.Corrida7, e.Corrida8, e.Burpee, e.Lunge, e.Farmer_Walk,\
        e.Abdominal, e.Jump_Squat, e.Medball_Alto, e.Medball_Solo, e.Terra_Remada
        from Experimentos e
        join Pessoas p ON e.ID_Pessoa = p.ID_Pessoa
        WHERE p.Nome = ? AND e.Data = ?
        ''', (nome,data))

        # Obter o resultado da consulta
        resultados = cursor.fetchone()

        # Exibir o resultado
        if resultados:
            #Data, Corrida, Burpee, Lunge, Farmer_Walk, Abdominal, Jump_Squat, Medball_Alto,\
            #Medball_Solo, Terra_Remada = resultados
            print(f"Dado encontrado para {nome} na data {data} com sucesso.")
            resultados = resultados[1:len(resultados)]

            # Abrindo o arquivo modelo
            doc = Document('Relatorio_dia_modelo.docx')

            # Atualizando o nome do aprticipante no relatório modelo
            for paragraph in doc.paragraphs:
                if("nome.do.participante" in paragraph.text):
                    paragraph.text = paragraph.text.replace("nome.do.participante",nome)

            # Atualizando a data do treino no relatório modelo
            for paragraph in doc.paragraphs:
                if("data.do.relatorio" in paragraph.text):
                    paragraph.text = paragraph.text.replace("data.do.relatorio",data)

            # Atualizando o plot de corrida
            rotulos = [f"Corrida {i+1}" for i in range(8)] # Gera rótulos "Corrida 1", "Corrida 2", ...

            # Configurar o gráfico
            plt.figure(figsize=(10, 6))  # Define o tamanho do gráfico
            tempos = resultados[0:8]
            plt.bar(rotulos, tempos, color='skyblue')  # Gráfico de barras com cor personalizada

            # Adicionar título e rótulos dos eixos
            plt.title("Medidas de Tempo para 8 Corridas")
            plt.xlabel("Corridas")
            plt.ylabel("Tempo (s)")
            plt.savefig('grafico_temp.png')

            palavra_chave = "grafico.corrida"  # Palavra-chave que será substituída pela imagem

            # Iterar sobre os parágrafos para localizar a palavra-chave
            for paragraph in doc.paragraphs:
                if palavra_chave in paragraph.text:
                    # Limpar o conteúdo do parágrafo e adicionar a imagem
                    paragraph.clear()  # Remove o texto da palavra-chave
                    paragraph.add_run().add_picture('grafico_temp.png', width=Inches(3))



            # Atualizando cada um dos tempos
            exercicios = ["tempo.corrida1", "tempo.corrida2", "tempo.corrida3", "tempo.corrida4", "tempo.corrida5", \
            "tempo.corrida6", "tempo.corrida7", "tempo.corrida8","tempo.burpee", "tempo.lunge", "tempo.farmer.walk", \
            "tempo.abdominal", "tempo.jump.squat", "tempo.medball.alto",\
            "tempo.medball.solo", "tempo.terra.remada"]
            for paragraph in doc.paragraphs:
                for tempo, exercicio in zip(resultados,exercicios):
                    if(exercicio in paragraph.text):
                        paragraph.text = paragraph.text.replace(exercicio,str(tempo))



            # Atualizando o tempo total
            tempo_total = sum(resultados)
            for paragraph in doc.paragraphs:
                if("tempo.total" in paragraph.text):
                    paragraph.text = paragraph.text.replace("tempo.total",str(tempo_total))

            # Converter a string em um objeto datetime
            data_objeto = datetime.strptime(data, "%d/%m/%Y")

            # Formatar o objeto datetime em uma nova string com o formato desejado
            data_formatada = data_objeto.strftime("%d-%m-%Y")

            # Salvando o nome arquivo .docx
            doc.save("Relatorios/" + nome + "_" + data_formatada + ".docx")

            print("Relatório concluído.")

        else:
            print(f"Nenhum dado encontrado para {nome} na data {data}.")



    def gerar_planilha(self):
        print("Gerando planilha...")

        data = self.ui.calendarWidget.selectedDate().toString('dd/MM/yyyy')

        # Cria uma nova planilha
        workbook = Workbook()
        sheet = workbook.active

        # Adiciona dados na planilha
        sheet['A1'] = 'Nome'
        sheet['B1'] = 'Burpee'
        sheet['C1'] = 'Lunge'
        sheet['D1'] = 'Farmer_Walk'
        sheet['E1'] = 'Abdominal'
        sheet['F1'] = 'Jump_Squat'
        sheet['G1'] = 'Medball_Alto'
        sheet['H1'] = 'Medball_Solo'
        sheet['I1'] = 'Terra + Remada'
        sheet['J1'] = 'Corrida1'
        sheet['K1'] = 'Corrida2'
        sheet['L1'] = 'Corrida3'
        sheet['M1'] = 'Corrida4'
        sheet['N1'] = 'Corrida5'
        sheet['O1'] = 'Corrida6'
        sheet['P1'] = 'Corrida7'
        sheet['Q1'] = 'Corrida8'
        sheet['R1'] = 'Corida Total'
        sheet['S1'] = 'Tempo Total'
        # Conectar ao banco de dados
        conexao = sqlite3.connect('Banco de dados.db')
        cursor = conexao.cursor()
        # Executar a consulta SQL para obter todos os dados
        cursor.execute("SELECT * FROM Experimentos WHERE Data = ?", (data,))

        # Obter os resultados (todos os dados de cada pessoa)
        resultados = cursor.fetchall()

        for i,linha in enumerate(resultados):
            corridas = 0
            tempo_total = 0
            for i2,valor in enumerate(linha[3:len(linha)]):
                sheet[chr(66+i2) + str(i + 2)] = str(valor)
                tempo_total = tempo_total + valor
                if(i2 >= 8):
                    corridas = corridas + valor
            sheet['R' + str(i + 2)] = str(corridas)
            sheet['S' + str(i + 2)] = tempo_total

            cursor.execute('''
            SELECT p.Nome from Pessoas p WHERE p.Id_Pessoa = ?
            ''', (linha[1],))

            # Obter o resultado da consulta
            resultado = cursor.fetchone()
            sheet['A' + str(i + 2)] = resultado[0]

        # Converter a string em um objeto datetime
        data_objeto = datetime.strptime(data, "%d/%m/%Y")

        # Formatar o objeto datetime em uma nova string com o formato desejado
        data_formatada = data_objeto.strftime("%d-%m-%Y")

        # Salva a planilha
        workbook.save('Planilhas/planilha_' + data_formatada + '.xlsx')

        print("Planilha concluída.")


    def process_serial_data(self,serial_port):
        data_dict = {}
        current_key = None

        while True:
            # Lê uma linha da porta serial e a decodifica
            line = serial_port.readline().decode('utf-8').strip()
            print(line)
            # Verifica se a linha é a palavra chave "FIM"
            if line == "FIM":
                break

            # Se o primeiro caractere é "X", definimos uma nova chave
            if line.startswith("X"):
                current_key = line[1:]  # Define a chave removendo o "X"
                data_dict[current_key] = []  # Inicializa uma lista para essa chave
            else:
                # Tenta dividir e converter a linha em uma tupla de inteiros
                try:
                    # Verifica se todos os elementos são numéricos
                    values = tuple(int(value) for value in line.split(',') if value.isdigit())

                    # Adiciona aos dados apenas se for uma linha de dados válida (completa)
                    if current_key is not None and len(values) > 1:
                        data_dict[current_key].append(values)
                    else:
                        print(f"Linha ignorada (formato incorreto ou incompleto): {line}")

                except ValueError:
                    print(f"Linha inválida ignorada: {line}")  # Ignora linhas que não podem ser convertidas

        return data_dict

    def solicitar_informacoes_totem(self):
        IDTotem = self.ui.spinBoxIDTotem.value()
        print("Solitando informações ao totem ",IDTotem,"...")

        # Aquisição dos dados

        if self.serial_port.is_open:
            mensagem = "T" + str(IDTotem) + "\n"
            self.serial_port.write(mensagem.encode())
            self.ui.Atualizacoes.setText("Tentando conectar ao toten...")
            time.sleep(1)
            self.log_serial_data_to_file('log_serial.txt')

            print("Solitação concluída")

            self.atualizar_lista_totens_recebidos()

    def atualizar_nomes_participantes(self):
        # Conectar ao banco de dados
        conexao = sqlite3.connect('Banco de dados.db')
        cursor = conexao.cursor()

        cursor.execute("SELECT Nome FROM Pessoas")

        # Obter o resultado da consulta
        nomes = cursor.fetchall()
        self.ui.ComboBoxNomes.clear()
        lista = []
        for nome in nomes:
            lista.append(nome[0])

        lista.sort()

        for nome in lista:
            self.ui.ComboBoxNomes.addItem(nome)

        # Fechar a conexão com o banco de dados
        conexao.close()



    def atualizar_lista_totens_recebidos(self):
        #Lendo as informações necessárias
        IDTotem = self.ui.spinBoxIDTotem.value() #ID do totem
        texto_atual = self.ui.Atualizacoes.text() #Texto atual que está impresso
        horario_atual = datetime.now().strftime("%H:%M:%S")
        nova_linha = "\n" + "Informações do totem " + str(IDTotem) + " recebida ás " + horario_atual
        print(texto_atual)
        novo_texto = texto_atual + nova_linha
        self.ui.Atualizacoes.setText(novo_texto)

    def Limpar_atualizacoes_totens(self):
        self.ui.Atualizacoes.setText("Atualizações dos totens:")

    def atualizar_portas(self):
        ports = serial.tools.list_ports.comports()
        self.ui.comboBoxUSB.clear()
        for port in ports:
            self.ui.comboBoxUSB.addItem(port[0])

    def conectar_usb(self):
        PortaCom = self.ui.comboBoxUSB.currentText()
        if(self.serial_port == None):
            self.serial_port = serial.Serial(PortaCom, 115200,timeout=None)
        if self.serial_port.is_open:
            if(self.ui.pushButtonUSB.text() == "Conectar"):
                print("Conectado")
                self.ui.pushButtonUSB.setText("Desconectar")
            else:
                print("Desconectado da Porta" + PortaCom)
                self.ui.pushButtonUSB.setText("Conectar")
                self.serial_port.close()
                self.serial_port = None

    def log_serial_data_to_file(self, output_file):
        # Abre o arquivo em modo de escrita
        with open(output_file, 'a') as file:
            while True:
                if self.serial_port.in_waiting > 0:
                    # Lê uma linha da porta serial
                    line = self.serial_port.readline().decode('utf-8').strip()
                    print(f"Recebido: {line}")

                    if(line == "FIM"):
                        break

                    if not all(char.isdigit() or char == ',' for char in line):
                                    continue

                    # Grava no arquivo
                    file.write(line + '\n')
                    file.flush()  # Garante que os dados sejam salvos imediatamente no disco

    def criar_experimentos(self):
        with open('Configuração dos IDs.txt', 'r') as arquivo:
            linhas = arquivo.readlines()
            for linha in linhas:
                self.experimentos.append(Experimento(linha.split(",")[0], int(linha.split(",")[1])))

    def criar_dicionario_exercicios(self):
        with open('Configuração dos totens.txt') as arquivo:
            linhas = arquivo.readlines()
            for linha in linhas:
                self.dicionario_exercicios[linha.split(",")[0]] = int(linha.split(",")[1])

if __name__ == "__main__":
    app = QApplication(sys.argv)
    widget = Widget()
    widget.show()
    sys.exit(app.exec())


